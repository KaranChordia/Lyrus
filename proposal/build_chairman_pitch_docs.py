from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "proposal"


def _set_page_margins(doc: Document, *, top=0.8, bottom=0.8, left=0.9, right=0.9) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)


def _style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(18)
    h1.font.bold = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(13)
    h2.font.bold = True


def _add_rule(paragraph, *, color_hex="0F172A") -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_callout(doc: Document, title: str, body: str, *, accent_rgb=(11, 79, 140)) -> None:
    # A light-weight "callout box" using a 1x1 table for consistent padding.
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    cell.width = Inches(6.2)

    # Cell shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")  # slate-100
    tcPr.append(shd)

    # Add a left border accent via paragraph rule + colored title.
    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(*accent_rgb)
    p1.paragraph_format.space_after = Pt(2)

    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)

    doc.add_paragraph("")  # spacing after callout


def _add_two_col_table(
    doc: Document, rows: list[tuple[str, str]], *, col1_width=2.2, col2_width=4.0
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    hdr = table.rows[0].cells
    hdr[0].text = "What We Focus On"
    hdr[1].text = "What It Means for Lyrus"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E2E8F0")  # slate-200
        tcPr.append(shd)
    hdr[0].width = Inches(col1_width)
    hdr[1].width = Inches(col2_width)

    for left, right in rows:
        row = table.add_row().cells
        row[0].width = Inches(col1_width)
        row[1].width = Inches(col2_width)
        row[0].text = left
        row[1].text = right
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    doc.add_paragraph("")


@dataclass(frozen=True)
class PreparedFor:
    name: str
    title: str
    company: str


def build_chairman_brief(out_path: Path) -> None:
    pf = PreparedFor(
        name="Mr. C.P. Bothra",
        title="Chairman",
        company="Lyrus Life Sciences",
    )
    today = date.today().strftime("%d %b %Y")

    doc = Document()
    _set_page_margins(doc)
    _style_base(doc)

    # Title block
    p = doc.add_paragraph("Chairman Brief", style="Heading 1")
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph("A practical growth + execution partnership (combined retainer)")
    p.paragraph_format.space_after = Pt(6)
    _add_rule(doc.add_paragraph(""), color_hex="CBD5E1")

    meta = doc.add_paragraph()
    meta.add_run("Prepared for: ").bold = True
    meta.add_run(f"{pf.name}, {pf.title}, {pf.company}\n")
    meta.add_run("Prepared by: ").bold = True
    meta.add_run("Karan Chordia (Draft)\n")
    meta.add_run("Date: ").bold = True
    meta.add_run(today)
    meta.paragraph_format.space_after = Pt(10)

    # One-line positioning
    _add_callout(
        doc,
        "The simple idea:",
        "Help Lyrus win more high-quality international business and execute it faster, "
        "with better documentation discipline, clearer milestone visibility, and stronger distributor-facing communication.",
    )

    doc.add_paragraph("What matters at Chairman level", style="Heading 2")
    rows = [
        (
            "Reduce concentration risk",
            "A stronger outbound presence that attracts more distributors / partners, so growth is not dependent on a single channel.",
        ),
        (
            "Scale execution speed without risking quality",
            "Make day-to-day work more repeatable with clear workflows, templates, and assistive systems that reduce rework and firefighting.",
        ),
        (
            "Strengthen regulated-market credibility",
            "Position Lyrus as a serious, high-trust partner for regulated markets with proof-driven communication and disciplined documentation.",
        ),
    ]
    _add_two_col_table(doc, rows)

    doc.add_paragraph("What I will deliver (combined retainer)", style="Heading 2")
    _add_callout(
        doc,
        "1) Content, website, and digital presence management",
        "A consistent communication engine: case studies, capability deck refresh, website creation/updates, distributor-ready brochures, "
        "and ongoing communication management (LinkedIn/newsletters/PR-style updates) with a regulated-market tone.",
        accent_rgb=(2, 132, 199),
    )
    _add_callout(
        doc,
        "2) Internal workflows and operating dashboards",
        "Simple, practical internal systems that reduce coordination overhead: milestone/stability tracking, decision logs, document status, "
        "handoffs, and partner (CMO) status tracking. Built to be used daily, not to impress in demos.",
        accent_rgb=(14, 116, 144),
    )
    _add_callout(
        doc,
        "3) AI agents that remove repetitive effort (controlled and auditable)",
        "A set of tightly-scoped AI assistants that support regulated workflows: drafting, checking, summarizing, and indexing. "
        "Deployed in a controlled manner with clear boundaries and human sign-off.",
        accent_rgb=(21, 128, 61),
    )

    doc.add_paragraph("High-impact AI agents and workflows (examples)", style="Heading 2")
    doc.add_paragraph(
        "These are designed to improve day-to-day operations. Each one is scoped so it can be piloted safely with anonymized or dummy data first."
    )
    agent_rows = [
        (
            "Regulatory / Dossier Drafting Assistant",
            "Turns structured lab outputs into first-draft dossier sections and responses, using your templates and checklists; humans review and approve.",
        ),
        (
            "QA Batch Record Checker",
            "Runs consistency checks on batch records, stability tables, and deviations; flags missing fields, out-of-range entries, and formatting errors before review.",
        ),
        (
            "Deviation + CAPA Summarizer",
            "Takes investigation notes and produces a concise, standard summary and action list aligned to your internal format.",
        ),
        (
            "Regulatory Intelligence Monitor",
            "Watches relevant updates and summarizes what changed, what it affects, and what actions are required, in a weekly brief.",
        ),
        (
            "Scientific Literature + Patent Scout",
            "Continuously scans public sources, summarizes what matters for your focus areas, and proposes shortlists for the team to evaluate.",
        ),
        (
            "Internal Knowledge Base Builder",
            "Organizes learnings from completed projects into a searchable, permissioned knowledge base so new teams do not repeat avoidable mistakes.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.0)
    table.rows[0].cells[0].text = "Agent / Workflow"
    table.rows[0].cells[1].text = "Operational Impact"
    for c in table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = c._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "E2E8F0")
        tcPr.append(shd)
    for k, v in agent_rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v
        r[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        r[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph("")

    doc.add_paragraph("How we start (low-risk)", style="Heading 2")
    doc.add_paragraph(
        "We start with a 30-day pilot, then convert into a retainer only if results and working style are satisfactory."
    )
    pilot_rows = [
        ("Week 1", "Alignment workshop + define 3–5 measurable targets + collect existing collateral and SOP templates."),
        ("Week 2", "Ship v1 of distributor-facing collateral (1 case study + 1 brochure/deck section) and a working internal workflow dashboard prototype."),
        ("Week 3", "Pilot 1 AI assistant on a narrow workflow (example: QA checker or dossier draft), using anonymized/dummy inputs first."),
        ("Week 4", "Finalize deliverables + handover + propose a 90-day retainer roadmap (content + website + workflows + agents) with scope and costs."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.2)
    table.columns[1].width = Inches(5.0)
    table.rows[0].cells[0].text = "Timeline"
    table.rows[0].cells[1].text = "What You Get"
    for c in table.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
        c._tc.get_or_add_tcPr().append(OxmlElement("w:tcBorders"))
    for k, v in pilot_rows:
        r = table.add_row().cells
        r[0].text = k
        r[1].text = v
        r[0].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        r[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    doc.add_paragraph("")

    doc.add_paragraph("What I need from Lyrus (minimal)", style="Heading 2")
    needs = [
        "One point-of-contact for BD collateral review (1 hour/week).",
        "One point-of-contact for execution workflows (R&D / QA / regulatory) (1 hour/week).",
        "A sample set of non-confidential documents (or anonymized samples) to build templates and demos.",
        "Approval to pilot internally first (no external exposure until you sign off).",
    ]
    for n in needs:
        p = doc.add_paragraph(f"• {n}")
        p.paragraph_format.space_after = Pt(1)
    doc.add_paragraph("")

    doc.add_paragraph("Decision ask", style="Heading 2")
    _add_callout(
        doc,
        "Ask:",
        "Approve a 30-day pilot. After the pilot, decide whether to continue as a combined retainer for "
        "business development content + execution visibility + documentation discipline.",
        accent_rgb=(124, 58, 237),
    )

    footer = doc.add_paragraph("Confidential: For discussion with Lyrus Life Sciences leadership only.")
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(71, 85, 105)  # slate-600

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_one_pager(out_path: Path) -> None:
    pf = PreparedFor(
        name="Mr. C.P. Bothra",
        title="Chairman",
        company="Lyrus Life Sciences",
    )
    today = date.today().strftime("%d %b %Y")

    doc = Document()
    _set_page_margins(doc, top=0.7, bottom=0.7, left=0.85, right=0.85)
    _style_base(doc)

    p = doc.add_paragraph("One-Page Summary", style="Heading 1")
    p.paragraph_format.space_after = Pt(0)
    p2 = doc.add_paragraph("Growth + Execution Partnership (Combined Retainer)")
    p2.paragraph_format.space_after = Pt(6)
    _add_rule(doc.add_paragraph(""), color_hex="CBD5E1")

    meta = doc.add_paragraph()
    meta.add_run("Prepared for: ").bold = True
    meta.add_run(f"{pf.name}, {pf.title}, {pf.company}    ")
    meta.add_run("Date: ").bold = True
    meta.add_run(today)
    meta.paragraph_format.space_after = Pt(8)

    doc.add_paragraph("3 outcomes we optimize", style="Heading 2")
    outcomes = [
        "More diversified inbound opportunities from international distributors/partners.",
        "Faster, more repeatable day-to-day execution with less rework (workflows, templates, dashboards).",
        "Higher regulated-market credibility through proof-driven communication and disciplined documentation.",
    ]
    for o in outcomes:
        doc.add_paragraph(f"• {o}")

    doc.add_paragraph("")
    doc.add_paragraph("30-day pilot deliverables", style="Heading 2")
    pilot = [
        "1 distributor-facing case study (high credibility, regulated-market tone).",
        "Website + digital presence improvements (copy updates and a publishable content rhythm).",
        "A working internal workflow dashboard prototype (milestones/stability/document status).",
        "One narrow AI assistant pilot (example: QA checker or dossier drafting assistant) with clear boundaries and human sign-off.",
    ]
    for x in pilot:
        doc.add_paragraph(f"• {x}")

    doc.add_paragraph("")
    doc.add_paragraph("Decision ask", style="Heading 2")
    _add_callout(
        doc,
        "Approve the pilot.",
        "If results are strong, convert to a retainer that continues content + tools + workflow discipline as an ongoing capability.",
        accent_rgb=(124, 58, 237),
    )

    footer = doc.add_paragraph("Confidential: For discussion with Lyrus Life Sciences leadership only.")
    footer.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(71, 85, 105)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> None:
    build_chairman_brief(OUT_DIR / "Chairman_Brief_CP_Bothra.docx")
    build_one_pager(OUT_DIR / "One_Page_Summary_CP_Bothra.docx")
    print("OK")


if __name__ == "__main__":
    main()
